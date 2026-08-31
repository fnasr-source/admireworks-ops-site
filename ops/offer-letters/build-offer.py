#!/usr/bin/env python3
"""Build an Admireworks offer letter .docx FROM SCRATCH.

Rebuilt 2026-08-08 from the real letters in Drive (see letter-structure.md).

WHY FROM SCRATCH: the previous version required an owner-supplied
`template.docx` and hard-failed without it, so the generator could never run.
Every fixed string in the real letter is now in this file verbatim, so there is
no template to lose. The trade is that the logo image and exact fonts from the
Word original are not reproduced; brand typography is applied from the AW design
tokens instead (see BRAND below).

ENTITY SAFETY: the letterhead here is Admireworks Advertising LLC, taken from
the 2026 letters. The Drive file NAMED "[Offer Letter template]" carries the
legacy ADMIRE8 MARKETING MANAGEMENT L.L.C letterhead; building from that file
emits offers under the wrong legal entity. Do not "restore" it.

Usage:
  python3 build-offer.py --config offer.json [--output ~/Desktop/]
  python3 build-offer.py --print-schema

Requires: python-docx  (pip3 install python-docx)
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from datetime import date, datetime
from pathlib import Path

HERE = Path(__file__).resolve().parent

# ── Brand (CLAUDE.md design tokens) ───────────────────────────────────────────
BRAND = {
    "navy": "001A70",
    "gold": "CC9F53",
    # Jaymont (headline) and Akkurat Pro (body) are the brand faces but are not
    # installed on most machines; Word silently substitutes when a face is
    # missing, so we name a serif/sans pair that degrades predictably.
    "body_font": "Calibri",
    "arabic_font": "Arial",
}

ENTITY_AR = [
    "ادميري وركس للإعلان (ش.ذ.م.م)",
    "أم هرير ٢، شارع مستشفى راشد",
    "جلف تاورز، جناح A2 مكتب ٥٢٢",
    "ص.ب.، ١٢٥٠٣٣ ، دبي",
    "الإمارات العربية المتحدة",
]
ENTITY_EN = [
    "Admireworks Advertising LLC",
    "Umm hurair 2, Rashid Hospital Street",
    "Gulf Towers, A2, Office 522",
    "PO Box 125033, Dubai",
    "United Arab Emirates",
    "admireworks.com",
]

OPENING = ("We are pleased to offer you the position of {role_title} with Admireworks, "
           "you'll be expected to report directly to our {reports_to_title}, Mr. Fouad Nasseredin.")

CLOSING = [
    "This offer is contingent upon a satisfactory background check and reference check.",
    "If you choose to accept this offer, please sign and date this letter and return it to me at your earliest convenience.",
    "We are excited about the potential of you joining our team and believe that Admireworks will provide a professional opportunity that is in line with your skills and expertise.",
    "Thank you for considering this offer. We look forward to your positive response.",
]
SIGNATURE = ["Best Regards,", "Fouad Nasseredin", "Managing Partner Signature / Date", "Admireworks"]

SCHEMA = {
    "candidate_name": "Sally Tarek           # full name, used in the filename",
    "first_name": "Sally                     # optional; derived from candidate_name",
    "role_title": "Account Manager",
    "role_family": "account-manager          # picks duty-libraries/{family}.md",
    "letter_date": "2026-05-16               # ISO; rendered as '16 May 2026'",
    "start_date": "2026-06-01                # ISO; rendered as '1st of June 2026'",
    "base_salary": "20,000 EGP per month",
    "reports_to_title": "Managing Partner    # optional, defaults to Managing Partner",
    "engagement_clause": "(optional) verbatim clause; see engagement-templates.md",
    "bonus_block": "(optional) verbatim block; see bonus-structures.md",
    "duties_override": "(optional) list of {heading, bullets[]} to bypass the library",
}


def die(msg: str) -> None:
    print(f"ERROR: {msg}", file=sys.stderr)
    sys.exit(1)


# ── Date rendering: the letter uses TWO formats ───────────────────────────────
def fmt_letter_date(value: str) -> str:
    """Top-of-letter date: bare numeral, e.g. '16 May 2026'."""
    d = _parse(value)
    return f"{d.day} {d.strftime('%B %Y')}" if d else value


def fmt_start_date(value: str) -> str:
    """Start date: ordinal, e.g. '1st of June 2026'."""
    d = _parse(value)
    if not d:
        return value
    n = d.day
    suffix = "th" if 11 <= n % 100 <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")
    return f"{n}{suffix} of {d.strftime('%B %Y')}"


def _parse(value: str):
    for fmt in ("%Y-%m-%d", "%d %B %Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(str(value).strip(), fmt).date()
        except (ValueError, TypeError):
            continue
    return None


# ── Duty library ──────────────────────────────────────────────────────────────
def load_duties(role_family: str):
    path = HERE / "duty-libraries" / f"{role_family}.md"
    if not path.exists():
        die(f"no duty library for '{role_family}' ({path}). Derive one from the closest "
            "existing library plus the offer email's 'What you'll own' bullets, get owner "
            "approval, and save it there.")
    text = path.read_text(encoding="utf-8")
    if "STUB (per the aw-offer-letter skill" in text:
        die(f"duty library '{role_family}' is still a STUB. Derive the full 9-section "
            "library before building an offer with it.")
    sections, current = [], None
    for line in text.splitlines():
        s = line.strip()
        m = re.match(r"^(\d+)\.\s+(.*)$", s)
        if m:
            current = {"heading": f"{m.group(1)}. {m.group(2)}", "bullets": []}
            sections.append(current)
        elif current is not None and s and not s.startswith("#") and not s.startswith(">"):
            current["bullets"].append(s.lstrip("- ").strip())
    if len(sections) < 8:
        die(f"duty library '{role_family}' parsed {len(sections)} sections; expected 9.")
    return sections


# ── Document ──────────────────────────────────────────────────────────────────
def build(cfg: dict, out_dir: Path) -> Path:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        die("python-docx is required: pip3 install python-docx")

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BRAND["body_font"]
    normal.font.size = Pt(10.5)

    def para(text="", *, bold=False, size=None, color=None, align=None, space_after=6):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            if size:
                run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)
        return p

    # Letterhead: a real 2-column centred table, Arabic left / English right.
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ar_cell, en_cell = table.rows[0].cells
    for cell, lines, font in ((ar_cell, ENTITY_AR, BRAND["arabic_font"]),
                              (en_cell, ENTITY_EN, BRAND["body_font"])):
        cell.text = ""
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = font
            run.font.size = Pt(9)
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(BRAND["navy"])

    para(space_after=12)
    para(fmt_letter_date(cfg["letter_date"]))
    first = cfg.get("first_name") or str(cfg["candidate_name"]).split()[0]
    para(f"Dear {first},", space_after=12)

    para(OPENING.format(role_title=cfg["role_title"],
                        reports_to_title=cfg.get("reports_to_title", "Managing Partner")))
    if cfg.get("engagement_clause"):
        para(cfg["engagement_clause"])

    para("Compensation", bold=True, size=12, color=BRAND["navy"], space_after=4)
    bonus = cfg.get("bonus_block")
    if bonus:
        para(f"The starting salary for this position will be {cfg['base_salary']}, "
             "plus a performance bonus tied directly to client outcomes:")
        for line in str(bonus).splitlines():
            if line.strip():
                para(line.rstrip(), space_after=2)
            else:
                para(space_after=2)
    else:
        para(f"The starting salary for this position will be {cfg['base_salary']}.")

    para("Start date", bold=True, size=12, color=BRAND["navy"], space_after=4)
    para(f"Your official date of hire is {fmt_start_date(cfg['start_date'])}.", space_after=12)

    para("JOB DUTIES AND RESPONSIBILITIES", bold=True, size=12, color=BRAND["navy"], space_after=4)
    para(f"The {cfg['role_title']} will be responsible for the following:", space_after=8)

    duties = cfg.get("duties_override") or load_duties(cfg["role_family"])
    for section in duties:
        para(section["heading"], bold=True, space_after=2)
        for bullet in section["bullets"]:
            para(bullet, space_after=2)
        para(space_after=4)

    for line in CLOSING:
        para(line)
    para(space_after=10)
    for line in SIGNATURE:
        para(line, bold=(line == "Fouad Nasseredin"), space_after=2)

    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / f"{cfg['role_title']} - Job Offer - {cfg['candidate_name']}.docx"
    doc.save(str(out_path))

    archive = HERE / "past-offers"
    archive.mkdir(exist_ok=True)
    slug = re.sub(r"[^a-z0-9]+", "-", str(cfg["candidate_name"]).lower()).strip("-")
    d = _parse(cfg["letter_date"]) or date.today()
    shutil_copy(out_path, archive / f"{d.isoformat()}_{slug}_{cfg['role_family']}.docx")
    return out_path


def shutil_copy(src: Path, dst: Path) -> None:
    import shutil
    shutil.copyfile(src, dst)


def main() -> None:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--config", help="path to the offer config JSON")
    ap.add_argument("--output", default=str(Path.home() / "Desktop"))
    ap.add_argument("--print-schema", action="store_true", help="print the config shape and exit")
    args = ap.parse_args()

    if args.print_schema:
        print(json.dumps(SCHEMA, indent=2, ensure_ascii=False))
        return
    if not args.config:
        die("--config is required (or use --print-schema)")

    cfg = json.loads(Path(args.config).expanduser().read_text(encoding="utf-8"))
    for field in ("candidate_name", "role_title", "role_family", "letter_date", "start_date", "base_salary"):
        if not cfg.get(field):
            die(f"config missing required field '{field}'")

    out = build(cfg, Path(args.output).expanduser())
    print(f"OK: {out}")
    print("Archived to ops/offer-letters/past-offers/")
    print("Next: attach it to the cover email in ops/offer-letters/cover-email.md (DRAFT only, never auto-send).")


if __name__ == "__main__":
    main()
